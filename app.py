import streamlit as st
import io
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime, date

# ─── Page Config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Legal Note File Generator",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── Custom CSS ─────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

.stApp {
    background: linear-gradient(135deg, #0f0c29 0%, #1a1a4e 40%, #24243e 100%);
    min-height: 100vh;
}

/* Hero header */
.hero-header {
    background: linear-gradient(135deg, rgba(255,255,255,0.08) 0%, rgba(255,255,255,0.03) 100%);
    border: 1px solid rgba(255,255,255,0.12);
    border-radius: 20px;
    padding: 2.5rem 3rem;
    margin-bottom: 2rem;
    backdrop-filter: blur(20px);
    text-align: center;
}
.hero-title {
    font-family: 'Playfair Display', serif;
    font-size: 2.6rem;
    font-weight: 700;
    background: linear-gradient(135deg, #f8f0ff 0%, #c084fc 50%, #818cf8 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin: 0 0 0.5rem 0;
}
.hero-sub {
    color: rgba(255,255,255,0.55);
    font-size: 1rem;
    font-weight: 400;
    letter-spacing: 0.02em;
}

/* Section cards */
.section-card {
    background: rgba(255,255,255,0.05);
    border: 1px solid rgba(255,255,255,0.10);
    border-radius: 16px;
    padding: 1.8rem 2rem;
    margin-bottom: 1.5rem;
    backdrop-filter: blur(12px);
}
.section-title {
    font-size: 1.1rem;
    font-weight: 600;
    color: #c084fc;
    letter-spacing: 0.05em;
    text-transform: uppercase;
    margin-bottom: 1.2rem;
    padding-bottom: 0.6rem;
    border-bottom: 1px solid rgba(192,132,252,0.25);
}

/* Party cards */
.party-card {
    background: rgba(255,255,255,0.03);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 12px;
    padding: 1.2rem 1.5rem;
    margin-bottom: 1rem;
}
.party-label {
    font-size: 0.85rem;
    font-weight: 600;
    color: #818cf8;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    margin-bottom: 0.8rem;
}

/* Inputs */
div[data-baseweb="input"] > div,
div[data-baseweb="textarea"] > div,
div[data-baseweb="select"] > div {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: 10px !important;
    color: #f1f5f9 !important;
}
div[data-baseweb="input"] > div:focus-within,
div[data-baseweb="textarea"] > div:focus-within {
    border-color: #c084fc !important;
    box-shadow: 0 0 0 3px rgba(192,132,252,0.2) !important;
}
.stTextInput label, .stTextArea label, .stSelectbox label,
.stDateInput label, .stNumberInput label {
    color: rgba(255,255,255,0.75) !important;
    font-size: 0.875rem !important;
    font-weight: 500 !important;
}
input, textarea {
    color: #f1f5f9 !important;
}

/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, #7c3aed 0%, #4f46e5 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.6rem 1.6rem !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    letter-spacing: 0.02em !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 15px rgba(124,58,237,0.35) !important;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(124,58,237,0.5) !important;
}
.generate-btn > button {
    background: linear-gradient(135deg, #059669 0%, #0d9488 100%) !important;
    box-shadow: 0 4px 15px rgba(5,150,105,0.35) !important;
    font-size: 1rem !important;
    padding: 0.8rem 2.5rem !important;
    width: 100% !important;
}
.generate-btn > button:hover {
    box-shadow: 0 6px 20px rgba(5,150,105,0.5) !important;
}

/* Download button */
.stDownloadButton > button {
    background: linear-gradient(135deg, #0284c7 0%, #0369a1 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.8rem 2rem !important;
    font-weight: 600 !important;
    font-size: 1rem !important;
    width: 100% !important;
    box-shadow: 0 4px 15px rgba(2,132,199,0.35) !important;
}
.stDownloadButton > button:hover {
    box-shadow: 0 6px 20px rgba(2,132,199,0.5) !important;
}

/* Dividers */
hr { border-color: rgba(255,255,255,0.08) !important; }

/* Info boxes */
.stInfo {
    background: rgba(129,140,248,0.1) !important;
    border: 1px solid rgba(129,140,248,0.3) !important;
    border-radius: 10px !important;
    color: #c7d2fe !important;
}
.stSuccess {
    background: rgba(5,150,105,0.1) !important;
    border: 1px solid rgba(5,150,105,0.3) !important;
    border-radius: 10px !important;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background: rgba(15,12,41,0.95) !important;
    border-right: 1px solid rgba(255,255,255,0.08) !important;
}

.badge {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 600;
    margin-left: 8px;
    vertical-align: middle;
}
.badge-c { background: rgba(239,68,68,0.2); color: #f87171; border: 1px solid rgba(239,68,68,0.3); }
.badge-d { background: rgba(59,130,246,0.2); color: #93c5fd; border: 1px solid rgba(59,130,246,0.3); }
</style>
""", unsafe_allow_html=True)

# ─── CONSTANTS / DEFAULTS ────────────────────────────────────────────────────
TEMPLATE_PATH = r"c:\Users\vpath\OneDrive\Desktop\SUN APP\L KAMSAMMA   3263.docx"

DEFAULTS = {
    "file_number": "G/3263/2025",
    "division": "Kandukur Division",
    "district": "Ranga Reddy District",
    "act_year": "2007",
    "mandal": "Maheshwaram Mandal",
    "application_date": "09.07.2025",
    "hearing_date": "07.11.2025",
    "hearing_time": "11:00 AM",
    "hearing_month": "11",
    "hearing_year": "2025",
    # Petitioner
    "petitioner_name": "Smt. Lingala Kamsamma",
    "petitioner_relation": "W/o Late L. Narsimha Das",
    "petitioner_age": "60",
    "petitioner_hno": "1-42",
    "petitioner_street": "",
    "petitioner_village": "Ravirala Village",
    "petitioner_mandal": "Maheshwaram Mandal",
    "petitioner_district": "Ranga Reddy District",
    "petitioner_pin": "",
    "petitioner_mobile": "",
}

RESPONDENT_DEFAULTS = [
    {
        "name": "Smt. Bhargavi",
        "relation": "",
        "form": "Form C",
        "hno": "",
        "street": "",
        "village": "",
        "mandal": "",
        "district": "",
        "pin": "",
        "mobile": "",
    },
    {
        "name": "Smt. Ashwini",
        "relation": "",
        "form": "Form C",
        "hno": "",
        "street": "",
        "village": "",
        "mandal": "",
        "district": "",
        "pin": "",
        "mobile": "",
    },
]

# ─── Session State Init ──────────────────────────────────────────────────────
# Initialize all form field defaults into session_state so widgets are pre-filled
for _k, _v in DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

if "respondents" not in st.session_state:
    st.session_state.respondents = [dict(r) for r in RESPONDENT_DEFAULTS]
if "generated_doc" not in st.session_state:
    st.session_state.generated_doc = None

# ─── Helper Functions ─────────────────────────────────────────────────────────
def build_address(hno, street, village, mandal, district, pin):
    """Compose address string from parts."""
    parts = []
    if hno:
        parts.append(f"H.No. {hno}")
    if street:
        parts.append(street)
    if village:
        parts.append(village)
    if mandal:
        parts.append(mandal)
    if district:
        parts.append(district)
    if pin:
        parts.append(f"PIN - {pin}")
    return ", ".join(parts)

def build_petitioner_full(d):
    name = d["petitioner_name"]
    rel = d["petitioner_relation"]
    age = d["petitioner_age"]
    addr = build_address(
        d["petitioner_hno"], d["petitioner_street"],
        d["petitioner_village"], d["petitioner_mandal"],
        d["petitioner_district"], d["petitioner_pin"]
    )
    mobile = d.get("petitioner_mobile", "")
    parts = [name]
    if rel:
        parts.append(rel)
    if age:
        parts.append(f"aged about {age} years")
    if addr:
        parts.append(f"and resident of {addr}")
    if mobile:
        parts.append(f"Mobile: {mobile}")
    return ", ".join(parts) if len(parts) <= 1 else parts[0] + ", " + ", ".join(parts[1:])

def build_respondent_full(r):
    name = r["name"]
    rel = r.get("relation", "")
    addr = build_address(
        r.get("hno", ""), r.get("street", ""),
        r.get("village", ""), r.get("mandal", ""),
        r.get("district", ""), r.get("pin", "")
    )
    mobile = r.get("mobile", "")
    parts = [name]
    if rel:
        parts.append(rel)
    if addr:
        parts.append(f"resident of {addr}")
    if mobile:
        parts.append(f"Mobile: {mobile}")
    return ", ".join(parts)

def replace_in_para(para, mapping):
    """Replace placeholder keys with values in a paragraph, preserving run formatting.
    Guards against paragraphs with no runs to prevent document corruption.
    """
    if not para.runs:          # skip empty / image-only paragraphs
        return
    full_text = "".join(run.text for run in para.runs)
    new_text = full_text
    for key, val in mapping.items():
        new_text = new_text.replace(key, str(val))
    if new_text == full_text:
        return
    # Put all text into first run, blank out the rest (preserves first run's formatting)
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def generate_document(d, respondents):
    """Build the output .docx from scratch using the template structure."""
    doc = Document(TEMPLATE_PATH)

    petitioner_full = build_petitioner_full(d)
    petitioner_name = d["petitioner_name"]
    pet_addr = build_address(
        d["petitioner_hno"], d["petitioner_street"],
        d["petitioner_village"], d["petitioner_mandal"],
        d["petitioner_district"], d["petitioner_pin"]
    )

    # Respondents split by form type
    form_c_resp = [r for r in respondents if r["form"] == "Form C"]
    form_d_resp = [r for r in respondents if r["form"] == "Form D"]

    all_resp_names = ", ".join(r["name"] for r in respondents)
    form_c_names = ", ".join(r["name"] for r in form_c_resp)
    form_d_names = ", ".join(r["name"] for r in form_d_resp)

    file_no = d["file_number"]
    division = d["division"]
    district = d["district"]
    mandal = d["mandal"]
    app_date = d["application_date"]
    hearing_date = d["hearing_date"]
    hearing_time = d["hearing_time"]
    hearing_month = d["hearing_month"]
    hearing_year = d["hearing_year"]

    # Build paragraph-level mapping
    mapping = {
        "G/3263/2025": file_no,
        "Smt. Lingala Kamsamma": petitioner_name,
        "Lingala Kamsamma, W/o Late L. Narsimha Das, aged about 60 years and resident of H.No. 1-42, Ravirala Village, Maheshwaram Mandal, Ranga Reddy District": petitioner_full,
        "Lingala Kamsamma, W/o Late L. Narsimha Das, aged about 60 years, and resident of H.No. 1-42, Ravirala Village, Maheshwaram Mandal, Ranga Reddy District": petitioner_full,
        "09.07.2025": app_date,
        "07.11.2025": hearing_date,
        "11:00 AM": hearing_time,
        "Kandukur Division": division,
        "Ranga Reddy District": district,
        "Maheshwaram Mandal": mandal,
        "Smt. Bhargavi and Smt. Ashwini": all_resp_names,
        "   .11.2025": f"  .{hearing_month}.{hearing_year}",
        "W/o Late L. Narsimha Das, R/o Ravirala Village, Maheshwaram Mandal, Ranga Reddy District": (
            f"{d['petitioner_relation']}, R/o {pet_addr}" if d['petitioner_relation'] else f"R/o {pet_addr}"
        ),
    }

    # Process all paragraphs
    for para in doc.paragraphs:
        replace_in_para(para, mapping)

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, mapping)

    # ── Now surgically update the respondent list paragraphs ──
    # In the original: P11 = 'Smt. Bhargavi', P12 = 'Smt. Ashwini'
    # We need to replace those two paragraphs with our respondent list (for Form C)

    # Find paragraphs that contained the respondent names originally
    # They appear right after "A notice in Form C (Notice to the Opposite Party) has also been prepared..."
    # We'll identify them by scanning
    resp_para_indices = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # These are short lines that are just respondent names (single names)
        if text in ["Smt. Bhargavi", "Smt. Ashwini", form_c_names]:
            resp_para_indices.append(i)

    # Update the "To" section in Form C (numbered respondents)
    # P44 and P45 are the numbered respondent addresses in Form C
    to_paras_formC = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("1 Smt.") or text.startswith("2. Smt.") or text.startswith(f"1 {petitioner_name}") or text.startswith(f"2. {petitioner_name}"):
            to_paras_formC.append(i)

    # Build new "To" content for Form C - addressed to opposite parties
    if form_c_resp:
        # Build the numbered list of Form C respondents (including mobile)
        form_c_to_lines = []
        for idx, r in enumerate(form_c_resp):
            addr = build_address(r.get("hno",""), r.get("street",""), r.get("village",""), r.get("mandal",""), r.get("district",""), r.get("pin",""))
            rel   = r.get("relation", "")
            mob   = r.get("mobile", "")
            line_parts = [r["name"]]
            if rel:
                line_parts.append(rel)
            if addr:
                line_parts.append(f"resident of {addr}")
            if mob:
                line_parts.append(f"Mobile: {mob}")
            num_prefix = f"{idx+1}{'.' if idx > 0 else ''} "
            form_c_to_lines.append(num_prefix + ", ".join(line_parts))

        # Update the numbered paragraphs if found
        if to_paras_formC and len(to_paras_formC) >= 1:
            for j, para_idx in enumerate(to_paras_formC):
                if j < len(form_c_to_lines):
                    doc.paragraphs[para_idx].runs[0].text = form_c_to_lines[j]
                    for run in doc.paragraphs[para_idx].runs[1:]:
                        run.text = ""
                else:
                    doc.paragraphs[para_idx].runs[0].text = ""
                    for run in doc.paragraphs[para_idx].runs[1:]:
                        run.text = ""

    # ── Save via a real temp file, then read back into BytesIO ──
    # Saving directly to BytesIO can produce a file Word cannot open.
    # Writing to disk first and reading back is always reliable.
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)   # close the raw file descriptor; python-docx will open it
    try:
        doc.save(tmp_path)
        with open(tmp_path, "rb") as f:
            buf = io.BytesIO(f.read())
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
    buf.seek(0)
    return buf


# ─── UI Layout ───────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero-header">
  <div class="hero-title">⚖️ Legal Note File Generator</div>
  <div class="hero-sub">Maintenance & Welfare of Parents and Senior Citizens Act, 2007 — Automated Notice Generation</div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# SECTION 1: CASE DETAILS
# ══════════════════════════════════════════════════════
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown('<div class="section-title">📋 Case / File Details</div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1:
    file_number = st.text_input("File / Case Number", value=DEFAULTS["file_number"], key="file_number")
with col2:
    division = st.text_input("Division", value=DEFAULTS["division"], key="division")
with col3:
    district = st.text_input("District", value=DEFAULTS["district"], key="district")

col4, col5 = st.columns(2)
with col4:
    mandal = st.text_input("Mandal", value=DEFAULTS["mandal"], key="mandal")
with col5:
    application_date = st.text_input("Application Date (DD.MM.YYYY)", value=DEFAULTS["application_date"], key="application_date")

col6, col7, col8 = st.columns(3)
with col6:
    hearing_date = st.text_input("Hearing Date (DD.MM.YYYY)", value=DEFAULTS["hearing_date"], key="hearing_date")
with col7:
    hearing_time = st.text_input("Hearing Time", value=DEFAULTS["hearing_time"], key="hearing_time")
with col8:
    hearing_month = st.text_input("Notice Month (for header date)", value=DEFAULTS["hearing_month"], key="hearing_month")

st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# SECTION 2: PETITIONER / APPLICANT DETAILS
# ══════════════════════════════════════════════════════
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown('<div class="section-title">👤 Petitioner / Applicant Details</div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1:
    petitioner_name = st.text_input("Full Name", value=DEFAULTS["petitioner_name"], key="petitioner_name")
with col2:
    petitioner_relation = st.text_input("Relation (e.g. W/o, S/o, D/o)", value=DEFAULTS["petitioner_relation"], key="petitioner_relation")
with col3:
    petitioner_age = st.text_input("Age (years)", value=DEFAULTS["petitioner_age"], key="petitioner_age")

st.markdown('<div class="party-label" style="margin-top:1rem;">📍 Petitioner Address</div>', unsafe_allow_html=True)
addr_col1, addr_col2, addr_col3 = st.columns(3)
with addr_col1:
    petitioner_hno = st.text_input("House Number", value=DEFAULTS["petitioner_hno"], key="petitioner_hno")
with addr_col2:
    petitioner_street = st.text_input("Street", value=DEFAULTS["petitioner_street"], key="petitioner_street")
with addr_col3:
    petitioner_village = st.text_input("Village", value=DEFAULTS["petitioner_village"], key="petitioner_village")

addr_col4, addr_col5, addr_col6, addr_col7 = st.columns(4)
with addr_col4:
    petitioner_mandal = st.text_input("Mandal", value=DEFAULTS["petitioner_mandal"], key="petitioner_mandal")
with addr_col5:
    petitioner_district = st.text_input("District", value=DEFAULTS["petitioner_district"], key="petitioner_district")
with addr_col6:
    petitioner_pin = st.text_input("PIN Code", value=DEFAULTS["petitioner_pin"], key="petitioner_pin")
with addr_col7:
    petitioner_mobile = st.text_input("Mobile Number", value=DEFAULTS["petitioner_mobile"], key="petitioner_mobile")

st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# SECTION 3: RESPONDENTS / OPPOSITE PARTIES
# ══════════════════════════════════════════════════════
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown("""
<div class="section-title">
  👥 Respondents / Opposite Parties
  &nbsp;
  <span style="font-size:0.75rem;font-weight:400;color:rgba(255,255,255,0.5);text-transform:none;letter-spacing:0;">
    Form C = Opposite Party Notice &nbsp;|&nbsp; Form D = Applicant Notice
  </span>
</div>
""", unsafe_allow_html=True)

# Buttons to add/remove respondents
btn_col1, btn_col2, _ = st.columns([1, 1, 4])
with btn_col1:
    if st.button("➕ Add Respondent"):
        st.session_state.respondents.append({
            "name": "", "relation": "", "form": "Form C",
            "hno": "", "street": "", "village": "",
            "mandal": "", "district": "", "pin": "", "mobile": ""
        })
        st.rerun()
with btn_col2:
    if len(st.session_state.respondents) > 1:
        if st.button("➖ Remove Last"):
            st.session_state.respondents.pop()
            st.rerun()

for i, resp in enumerate(st.session_state.respondents):
    st.markdown(f'<div class="party-card">', unsafe_allow_html=True)

    form_badge = "badge-c" if resp.get("form") == "Form C" else "badge-d"
    form_label = resp.get("form", "Form C")
    st.markdown(f'<div class="party-label">Respondent {i+1} <span class="badge {form_badge}">{form_label}</span></div>', unsafe_allow_html=True)

    r_col1, r_col2, r_col3 = st.columns(3)
    with r_col1:
        resp["name"] = st.text_input(f"Full Name", value=resp.get("name",""), key=f"resp_name_{i}")
    with r_col2:
        resp["relation"] = st.text_input(f"Relation (e.g. S/o, W/o)", value=resp.get("relation",""), key=f"resp_relation_{i}")
    with r_col3:
        form_options = ["Form C", "Form D"]
        current_form = resp.get("form", "Form C")
        resp["form"] = st.selectbox(
            f"Notice Form",
            options=form_options,
            index=form_options.index(current_form),
            key=f"resp_form_{i}",
            help="Form C = Opposite Party | Form D = Applicant"
        )

    st.markdown(f'<div class="party-label" style="font-size:0.78rem;color:rgba(255,255,255,0.45);margin-top:0.5rem;">Address</div>', unsafe_allow_html=True)
    a1, a2, a3 = st.columns(3)
    with a1:
        resp["hno"] = st.text_input("House Number", value=resp.get("hno",""), key=f"resp_hno_{i}")
    with a2:
        resp["street"] = st.text_input("Street", value=resp.get("street",""), key=f"resp_street_{i}")
    with a3:
        resp["village"] = st.text_input("Village", value=resp.get("village",""), key=f"resp_village_{i}")

    a4, a5, a6, a7 = st.columns(4)
    with a4:
        resp["mandal"] = st.text_input("Mandal", value=resp.get("mandal",""), key=f"resp_mandal_{i}")
    with a5:
        resp["district"] = st.text_input("District", value=resp.get("district",""), key=f"resp_district_{i}")
    with a6:
        resp["pin"] = st.text_input("PIN Code", value=resp.get("pin",""), key=f"resp_pin_{i}")
    with a7:
        resp["mobile"] = st.text_input("Mobile Number", value=resp.get("mobile",""), key=f"resp_mobile_{i}")

    st.markdown('</div>', unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# SECTION 4: PREVIEW SUMMARY
# ══════════════════════════════════════════════════════
with st.expander("📄 Preview Data Summary", expanded=False):
    st.markdown("**File Number:** " + st.session_state.get("file_number", file_number))
    st.markdown("**Division:** " + st.session_state.get("division", division))
    st.markdown("**Petitioner:** " + petitioner_name)
    pet_addr_preview = build_address(petitioner_hno, petitioner_street, petitioner_village, petitioner_mandal, petitioner_district, petitioner_pin)
    st.markdown("**Address:** " + pet_addr_preview)
    st.markdown("**Hearing:** " + hearing_date + " at " + hearing_time)
    st.markdown("---")
    for i, r in enumerate(st.session_state.respondents):
        badge = "🔴 Form C" if r.get("form") == "Form C" else "🔵 Form D"
        resp_addr = build_address(r.get("hno",""), r.get("street",""), r.get("village",""), r.get("mandal",""), r.get("district",""), r.get("pin",""))
        st.markdown(f"**Respondent {i+1}:** {r.get('name','')} ({r.get('relation','')}) — {badge}")
        if resp_addr:
            st.markdown(f"&nbsp;&nbsp;📍 {resp_addr}")

# ══════════════════════════════════════════════════════
# GENERATE BUTTON
# ══════════════════════════════════════════════════════
st.markdown("<br>", unsafe_allow_html=True)
gen_col1, gen_col2, gen_col3 = st.columns([1, 2, 1])
with gen_col2:
    st.markdown('<div class="generate-btn">', unsafe_allow_html=True)
    generate_clicked = st.button("📄 Generate Note File", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

if generate_clicked:
    # Gather all data
    data = {
        "file_number": file_number,
        "division": division,
        "district": district,
        "mandal": mandal,
        "application_date": application_date,
        "hearing_date": hearing_date,
        "hearing_time": hearing_time,
        "hearing_month": hearing_month,
        "hearing_year": hearing_date.split(".")[-1] if "." in hearing_date else "2025",
        "petitioner_name": petitioner_name,
        "petitioner_relation": petitioner_relation,
        "petitioner_age": petitioner_age,
        "petitioner_hno": petitioner_hno,
        "petitioner_street": petitioner_street,
        "petitioner_village": petitioner_village,
        "petitioner_mandal": petitioner_mandal,
        "petitioner_district": petitioner_district,
        "petitioner_pin": petitioner_pin,
        "petitioner_mobile": petitioner_mobile,
    }

    with st.spinner("⚙️ Generating document..."):
        try:
            buf = generate_document(data, st.session_state.respondents)
            st.session_state.generated_doc = buf.getvalue()
            st.success("✅ Document generated successfully!")
        except Exception as e:
            st.error(f"❌ Error generating document: {e}")
            import traceback
            st.code(traceback.format_exc())

if st.session_state.generated_doc:
    dl_col1, dl_col2, dl_col3 = st.columns([1, 2, 1])
    with dl_col2:
        filename = f"Note_File_{file_number.replace('/', '_')}.docx"
        st.download_button(
            label="⬇️ Download Note File (.docx)",
            data=st.session_state.generated_doc,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ── Footer ───────────────────────────────────────────────────────────────────
st.markdown("""
<div style="text-align:center;padding:2rem 0 1rem;color:rgba(255,255,255,0.25);font-size:0.8rem;">
  Legal Note File Generator &nbsp;·&nbsp; Maintenance & Welfare of Parents and Senior Citizens Act, 2007
</div>
""", unsafe_allow_html=True)
